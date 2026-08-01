"""
Universal Platform Scraper - WithOne.AI Knowledge Base
Scrapes actions/API calls from any platform on https://www.withone.ai/knowledge/
Supports: GitHub, Linear, Slack, Jira, Notion, GDrive, Asana, Monday, HubSpot,
Zendesk, Salesforce, Confluence, Figma, Google Calendar, PagerDuty, Datadog,
Discord, Zoom, Stripe, Intercom, Loom, GitHub Actions, Google Meet, Trello, etc.

FIXES:
- Heading detection rewritten: only promotes lines that are immediately followed
  by substantial content (not another short line), are Title Case / ALL CAPS,
  and never start with code-like characters.
- Content extraction: noise removal is tighter; code blocks are preserved intact.
- Table detection: generalised header pattern (any "X  Y  Z" 2-3 word columns).
- "Copy code" artefact skips the language tag BUT keeps the code that follows.
- Navigation uses direct URL instead of driver.back() to avoid stale-page issues.
"""

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from docx import Document
import re
import time
import sys
import os


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

LANGUAGE_TAGS = {
    "json", "bash", "shell", "html", "text", "yaml", "xml",
    "javascript", "python", "typescript", "ruby", "go", "java", "php",
}

# Characters that signal a line is *code / data*, not a heading
CODE_START_CHARS = (
    "{", "}", "[", "]", "-", "*", "|", "curl ", "#", "//", "/*",
    "GET ", "POST ", "PUT ", "PATCH ", "DELETE ", "HTTP/",
    '"', "'", "<", ">", "0", "1", "2", "3", "4", "5", "6", "7", "8", "9",
)

# Exact UI strings that should be stripped from scraped content
NOISE_EXACT: set[str] = {
    "Knowledge", "GitHub", "/", "|", "Get started", "Docs",
    "GitHub skills", "Copy code",
}

NOISE_CONTAINS: list[str] = [
    "Use this action in your agent",
    "Suggest edits",
    "Report an issue",
    "This knowledge powers One's agent infrastructure",
]


def is_code_like(text: str) -> bool:
    """Return True if the line looks like code / data rather than prose."""
    s = text.strip()
    if not s:
        return False
    if s.startswith(CODE_START_CHARS):
        return True
    # Contains assignment / key-value typical of JSON / YAML / code
    if re.search(r'[=:]\s*[\[{"\d]', s):
        return True
    # URL
    if re.match(r'https?://', s):
        return True
    return False


def is_heading_candidate(line: str, next_non_empty: str) -> bool:
    """
    A line is treated as a section heading if ALL of the following hold:
      1. Not empty.
      2. Short (≤ 70 chars, ≤ 8 words).
      3. Does NOT look like code / data.
      4. Does NOT end with sentence-terminating punctuation.
      5. Starts with an uppercase letter (Title Case or ALL CAPS).
      6. The very next non-empty line is longer / more content-rich,
         OR the next line itself looks like a heading of lower level
         (we only promote one level deep).
    """
    s = line.strip()
    if not s:
        return False
    if is_code_like(s):
        return False
    if len(s) > 70:
        return False
    words = s.split()
    if len(words) > 8:
        return False
    if s.endswith((".", ",", ";", ":", "?")):
        return False
    if not s[0].isupper():
        return False
    # Must be followed by something substantive
    if not next_non_empty:
        return False
    # Avoid promoting a line that looks exactly like the following line
    # (duplicate artefact from the scraper)
    if s.lower() == next_non_empty.strip().lower():
        return False
    return True


def is_list_item(text: str) -> bool:
    s = text.strip()
    if not s:
        return False
    if s.startswith(("- ", "* ", "• ")):
        return True
    # Numbered list: "1. " / "2) "
    if re.match(r'^\d+[.)]\s', s):
        return True
    return False


# ---------------------------------------------------------------------------
# Scraper
# ---------------------------------------------------------------------------

class PlatformScraper:
    def __init__(self, platform_name: str, platform_url: str):
        self.platform_name = platform_name
        self.platform_url = platform_url

        options = Options()
        options.add_argument("--no-sandbox")
        options.add_argument("--disable-dev-shm-usage")
        options.add_argument("start-maximized")

        self.driver = webdriver.Chrome(options=options)
        self.wait = WebDriverWait(self.driver, 15)
        self._list_page_url: str = platform_url  # updated after each page load

    # ------------------------------------------------------------------
    # Navigation helpers
    # ------------------------------------------------------------------

    def open_platform_page(self):
        print(f"Opening {self.platform_name} knowledge page...")
        self.driver.get(self.platform_url)
        time.sleep(3)
        self._list_page_url = self.driver.current_url
        print(f"✅ Page loaded\n")

    def get_total_pages(self) -> int:
        """Detect total number of pages via pagination buttons."""
        try:
            time.sleep(1)
            max_page = 1
            for btn in self.driver.find_elements(By.XPATH, "//button | //a"):
                try:
                    text = btn.text.strip()
                    if text and len(text) <= 3 and text.isdigit():
                        max_page = max(max_page, int(text))
                except Exception:
                    continue

            if max_page > 1:
                print(f"✅ Detected {max_page} pages for {self.platform_name}\n")
                return max_page

            print("ℹ️  Page count not detected. Using adaptive detection...\n")
            return 100
        except Exception as e:
            print(f"⚠️  Error detecting pages: {e}\n")
            return 50

    def go_to_next_page(self, current_page: int) -> bool:
        """Click pagination to advance one page."""
        next_page = current_page + 1
        selectors = [
            f"//button[normalize-space(text())='{next_page}'] | //a[normalize-space(text())='{next_page}']",
            "//button[contains(normalize-space(text()),'Next')] | //a[contains(normalize-space(text()),'Next')]",
            "//button[@aria-label='Next page'] | //a[@aria-label='Next page']",
            "//button[contains(@aria-label,'>')] | //a[contains(@aria-label,'>')]",
        ]
        for selector in selectors:
            try:
                btn = self.driver.find_element(By.XPATH, selector)
                if btn and btn.is_enabled():
                    btn.click()
                    time.sleep(2)
                    self._list_page_url = self.driver.current_url
                    return True
            except Exception:
                continue
        return False

    # ------------------------------------------------------------------
    # Content extraction
    # ------------------------------------------------------------------

    def _clean_raw_text(self, raw: str, action_title: str) -> str:
        """
        Strip navigation / UI chrome from scraped page text while preserving
        all meaningful content — including code blocks — in original order.
        """
        cleaned: list[str] = []
        skip_title_once = True          # drop the first occurrence of action_title
        title_norm = action_title.strip().lower()  # normalized version for fuzzy match

        for raw_line in raw.split("\n"):
            line = raw_line.rstrip()
            stripped = line.strip()

            # Always preserve blank lines (collapsed later)
            if not stripped:
                cleaned.append("")
                continue

            # Stop at knowledge metadata footer
            if stripped == "KNOWLEDGE METADATA":
                break

            # Drop exact UI noise tokens
            if stripped in NOISE_EXACT:
                continue

            # Drop the page title once (it's already used as the Word heading)
            # Use both exact match and normalized match to catch whitespace variants
            if skip_title_once and (stripped == action_title or stripped.lower() == title_norm):
                skip_title_once = False
                continue

            # Drop breadcrumb-style lines ending with /
            if re.match(r'^[\w\s]+/$', stripped):
                continue

            # Drop "Back to <platform>" nav link
            if re.match(r'^Back to ', stripped, re.IGNORECASE):
                continue

            # Drop noise substrings
            if any(token in stripped for token in NOISE_CONTAINS):
                continue

            cleaned.append(line)

        # Collapse runs of more than two consecutive blank lines
        deduped: list[str] = []
        blank_run = 0
        for line in cleaned:
            if line.strip() == "":
                blank_run += 1
                if blank_run <= 2:
                    deduped.append(line)
            else:
                blank_run = 0
                deduped.append(line)

        # Strip leading / trailing blank lines
        while deduped and deduped[0].strip() == "":
            deduped.pop(0)
        while deduped and deduped[-1].strip() == "":
            deduped.pop()

        return "\n".join(deduped)

    def extract_action_details(self, action_data: dict):
        """
        Navigate to the action's detail page, extract cleaned full-text content,
        then return directly to the saved list-page URL (avoids stale back-nav).
        """
        method    = action_data.get("method", "UNKNOWN")
        title     = action_data.get("title", "")
        endpoint  = action_data.get("endpoint", "")
        href      = action_data.get("href", "")

        if not href:
            print(f"  [!] No href for: {title}")
            return method, title, endpoint, ""

        # Build absolute URL
        if href.startswith("http"):
            detail_url = href
        elif href.startswith("/"):
            base = "/".join(self.driver.current_url.split("/")[:3])
            detail_url = base + href
        else:
            detail_url = self.driver.current_url.rsplit("/", 1)[0] + "/" + href

        try:
            self.driver.get(detail_url)
            time.sleep(1.5)

            # Prefer <main> / <article> over the whole body to avoid nav chrome
            try:
                container = self.driver.find_element(
                    By.XPATH, "//main | //article | //div[@role='main']"
                )
                raw_text = container.text
            except Exception:
                raw_text = self.driver.find_element(By.XPATH, "//body").text

            cleaned = self._clean_raw_text(raw_text, title)

        except Exception as e:
            print(f"  [!] Error fetching detail page for '{title}': {e}")
            cleaned = ""
        finally:
            # Always navigate back to the list page by URL — never driver.back()
            self.driver.get(self._list_page_url)
            time.sleep(1.5)

        return method, title, endpoint, cleaned

    # ------------------------------------------------------------------
    # Scraping loop
    # ------------------------------------------------------------------

    def scrape_actions(self, max_pages=None, max_actions=None):
        all_actions: list[str] = []

        if max_pages is None:
            max_pages = self.get_total_pages()

        current_page = 1

        while current_page <= max_pages:
            print(f"{'='*70}")
            print(f"PAGE {current_page}/{max_pages}")
            print(f"{'='*70}")

            try:
                time.sleep(1)

                # ---- collect card metadata (no clicks) ----
                action_links = self.driver.find_elements(
                    By.XPATH,
                    "//a[contains(@class,'group') and contains(@class,'flex')]",
                )
                print(f"[*] Found {len(action_links)} action elements on page {current_page}")

                action_data_list = []
                for link in action_links:
                    try:
                        try:
                            method_text = link.find_element(
                                By.XPATH, ".//span[contains(@class,'font-mono')]"
                            ).text.strip()
                        except Exception:
                            method_text = "?"

                        try:
                            title_text = link.find_element(
                                By.XPATH, ".//span[contains(@class,'text-sm')]"
                            ).text.strip()
                        except Exception:
                            title_text = "Untitled"

                        try:
                            endpoint_text = link.find_element(
                                By.XPATH, ".//p[contains(@class,'font-mono')]"
                            ).text.strip()
                        except Exception:
                            endpoint_text = ""

                        href = link.get_attribute("href") or ""
                        action_data_list.append(
                            {"method": method_text, "title": title_text,
                             "endpoint": endpoint_text, "href": href}
                        )
                    except Exception as e:
                        print(f"  [!] Error collecting card data: {e}")

                print(f"[*] Collected metadata for {len(action_data_list)} actions")
                print(f"[*] Extracting detailed content…\n")

                page_actions: list[str] = []

                for i, action_info in enumerate(action_data_list, 1):
                    if max_actions and len(all_actions) >= max_actions:
                        print(f"[*] Reached max_actions limit ({max_actions})")
                        return all_actions

                    method, title, endpoint, full_content = self.extract_action_details(action_info)

                    action_content = f"{method}\n{title}\n{endpoint}"
                    if full_content:
                        action_content += f"\n\n{full_content}"

                    if action_content and action_content not in all_actions:
                        all_actions.append(action_content)
                        page_actions.append(action_content)
                        print(f"  [{i:2d}] {method:6s} - {title[:55]}…")

                print(f"\n[+] Collected {len(page_actions)} actions on page {current_page}")
                print(f"[*] Total so far: {len(all_actions)}\n")

                if current_page < max_pages:
                    if self.go_to_next_page(current_page):
                        current_page += 1
                    else:
                        print("[OK] Reached last page\n")
                        break
                else:
                    break

            except Exception as e:
                print(f"[!] Error on page {current_page}: {e}\n")
                break

        return all_actions

    # ------------------------------------------------------------------
    # Word document output
    # ------------------------------------------------------------------

    def save_to_word(self, actions: list[str]) -> str:
        tools_folder = "Tools"
        os.makedirs(tools_folder, exist_ok=True)
        filename = os.path.join(tools_folder, f"{self.platform_name.lower()}-tools.docx")

        print(f"\n{'='*70}")
        print(f"Creating Word Document: {filename}")
        print(f"{'='*70}\n")

        doc = Document()

        for i, action in enumerate(actions, 1):
            parts = action.split("\n\n", 1)
            header_section  = parts[0] if parts else ""
            content_section = parts[1] if len(parts) > 1 else ""

            header_lines = header_section.split("\n")
            method   = header_lines[0].strip() if len(header_lines) > 0 else ""
            title    = header_lines[1].strip() if len(header_lines) > 1 else ""
            endpoint = header_lines[2].strip() if len(header_lines) > 2 else ""

            doc.add_heading(f"{i}. [{method}] {title}", level=2)

            ep_para = doc.add_paragraph(f"Endpoint: {endpoint}")
            ep_para.paragraph_format.left_indent = 200_000

            if content_section:
                self._add_formatted_content(doc, content_section)

            doc.add_paragraph()  # spacer

        try:
            doc.save(filename)
            saved_path = filename
        except PermissionError:
            import datetime
            stamp = datetime.datetime.now().strftime("%Y%m%d-%H%M%S")
            fallback = os.path.join(
                tools_folder, f"{self.platform_name.lower()}-tools-{stamp}.docx"
            )
            doc.save(fallback)
            saved_path = fallback
            print(f"[!] Target file locked. Saved to: {saved_path}")

        print(f"[OK] Saved {len(actions)} actions to: {saved_path}\n")
        return saved_path

    # ------------------------------------------------------------------
    # Document formatting helpers
    # ------------------------------------------------------------------

    def _add_formatted_content(self, doc: Document, content_text: str):
        """
        Walk the cleaned content line-by-line and write appropriately styled
        paragraphs / tables into the Word document.

        Key improvements over original:
        - Heading detection uses is_heading_candidate() — context-aware.
        - Code blocks (after "Copy code" artefact) are preserved verbatim in
          a monospace paragraph instead of being dropped.
        - Table detection is generalised (tab-separated OR pipe-separated).
        - "Copy code" + language tag lines are stripped but the code itself kept.
        """
        lines = content_text.split("\n")
        n = len(lines)

        def next_non_empty_line(from_idx: int) -> str:
            for j in range(from_idx + 1, n):
                if lines[j].strip():
                    return lines[j].strip()
            return ""

        i = 0
        while i < n:
            raw  = lines[i]
            stripped = raw.strip()

            # ---- blank line ----
            if stripped == "":
                doc.add_paragraph("")
                i += 1
                continue

            # ---- "Copy code" artefact: drop the tag (+ optional language tag),
            #      but keep the code that follows ----
            if stripped.lower().startswith("copy code"):
                i += 1
                # Skip a bare language-tag line if present
                if i < n and lines[i].strip().lower() in LANGUAGE_TAGS:
                    i += 1
                # The code itself follows — do NOT skip it; let the loop handle it
                continue

            # ---- tab-separated table ----
            if "\t" in raw:
                table_lines = []
                while i < n and "\t" in lines[i]:
                    table_lines.append(lines[i])
                    i += 1
                self._add_table_to_doc(doc, table_lines)
                continue

            # ---- pipe-separated table ----
            if stripped.startswith("|") and stripped.endswith("|"):
                table_lines = []
                while i < n:
                    s = lines[i].strip()
                    if s.startswith("|") and s.endswith("|"):
                        table_lines.append(lines[i])
                        i += 1
                    else:
                        break
                # Filter separator rows (---|---) before deciding
                data_rows = [
                    l for l in table_lines
                    if not re.match(r'^\|[-:\s|]+\|$', l.strip())
                ]
                if len(data_rows) >= 2:
                    self._add_table_to_doc(doc, data_rows, pipe_separated=True)
                else:
                    # Too short to be a real table — write as plain text
                    for l in table_lines:
                        p = doc.add_paragraph(l)
                        p.paragraph_format.left_indent = 400_000
                continue

            # ---- generalised parameter / field table heuristic ----
            # Detect a 2-3 column plain-text table whose first line looks like
            # column headers (two or three single-word, Title-Case tokens).
            col_header_match = re.match(
                r'^([A-Z][a-zA-Z]+)\s{2,}([A-Z][a-zA-Z]+)(\s{2,}([A-Z][a-zA-Z]+))?$',
                stripped
            )
            if col_header_match:
                # Peek ahead: if next lines also look like aligned columns, build a table
                col_headers = [g for g in col_header_match.groups() if g and g == g.strip()]
                table_rows  = [col_headers]
                i += 1
                while i < n:
                    candidate = lines[i].strip()
                    if not candidate:
                        break
                    if candidate.lower().startswith("copy code"):
                        break
                    # Split on 2+ spaces (column alignment)
                    parts = re.split(r'\s{2,}', candidate)
                    if 2 <= len(parts) <= 4:
                        table_rows.append(parts)
                        i += 1
                    else:
                        break

                if len(table_rows) >= 2:
                    # Convert to tab-separated and pass to table renderer
                    self._add_table_to_doc(
                        doc, ["\t".join(r) for r in table_rows]
                    )
                    continue
                else:
                    # Not a real table — fall through as a heading
                    i -= (len(table_rows) - 1)  # rewind

            # ---- heading ----
            nne = next_non_empty_line(i)
            if is_heading_candidate(stripped, nne):
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = 200_000
                para.add_run(stripped).bold = True
                i += 1
                continue

            # ---- list item ----
            if is_list_item(raw):
                para = doc.add_paragraph(style="List Bullet")
                para.paragraph_format.left_indent = 400_000
                text = stripped
                # Remove bullet prefix
                if text.startswith(("- ", "* ", "• ")):
                    text = text[2:].strip()
                elif re.match(r'^\d+[.)]\s', text):
                    text = re.sub(r'^\d+[.)]\s+', "", text)
                para.add_run(text)
                i += 1
                continue

            # ---- plain / code line ----
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = 400_000
            para.paragraph_format.word_wrap   = True
            # Use monospace font for code-like lines
            run = para.add_run(raw)
            if is_code_like(stripped):
                run.font.name = "Courier New"
                run.font.size = 180_000   # 9 pt

            i += 1

    def _add_table_to_doc(
        self,
        doc: Document,
        table_lines: list[str],
        pipe_separated: bool = False,
    ):
        """Render a list of row-strings as a Word table."""
        rows = []
        for line in table_lines:
            s = line.strip()
            if pipe_separated:
                if s.startswith("|"): s = s[1:]
                if s.endswith("|"):   s = s[:-1]
                cells = [c.strip() for c in s.split("|")]
            elif "\t" in s:
                cells = [c.strip() for c in s.split("\t")]
            else:
                cells = re.split(r'\s{2,}', s)

            # Skip separator rows (---  :---  etc.)
            if all(re.match(r'^[-:\s]*$', c) for c in cells):
                continue
            rows.append(cells)

        if not rows:
            return

        num_cols = max(len(r) for r in rows)
        if num_cols == 0:
            return

        # Pad rows to same width
        rows = [r + [""] * (num_cols - len(r)) for r in rows]

        table = doc.add_table(rows=len(rows), cols=num_cols)
        try:
            table.style = "Light Grid Accent 1"
        except Exception:
            pass

        for r_idx, row_data in enumerate(rows):
            for c_idx, cell_text in enumerate(row_data):
                cell = table.rows[r_idx].cells[c_idx]
                cell.text = cell_text
                if r_idx == 0:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True

        doc.add_paragraph()  # spacing after table

    # ------------------------------------------------------------------

    def close(self):
        self.driver.quit()
        print("[DONE] Browser closed\n")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main():
    print("\n" + "=" * 70)
    print("GITHUB API ACTION SCRAPER")
    print("=" * 70 + "\n")

    max_pages   = None
    max_actions = None

    for arg in sys.argv[1:]:
        if arg.startswith("--max-pages="):
            max_pages = int(arg.split("=")[1])
        elif arg.startswith("--max-actions="):
            max_actions = int(arg.split("=")[1])

    platform = "github"
    url      = "https://www.withone.ai/knowledge/github"

    scraper = None
    try:
        print(f"{'='*70}")
        print(f"SCRAPING: GITHUB")
        print(f"{'='*70}\n")

        scraper = PlatformScraper(platform, url)
        scraper.open_platform_page()

        print("Starting to collect actions from all pages…\n")
        actions = scraper.scrape_actions(max_pages=max_pages, max_actions=max_actions)

        if actions:
            print(f"\n[OK] Successfully collected {len(actions)} total actions")
            filename = scraper.save_to_word(actions)

            print(f"\n{'='*70}")
            print("SCRAPING COMPLETE")
            print(f"{'='*70}")
            print(f"[OK] github  —  {len(actions):5d} actions → {filename}")
            print(f"\n[OK] ALL DONE!\n")
        else:
            print("[!] No actions were collected")

    except Exception as e:
        print(f"\n[!] Error scraping GitHub: {e}")
    finally:
        if scraper:
            scraper.close()


if __name__ == "__main__":
    main()